from __future__ import annotations

import base64
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from convert import (  # noqa: E402
    ALIASES,
    Block,
    THEMES,
    canonical_theme,
    parse_docx,
    parse_markdown,
    render_html,
    _normalize_docx_blocks,
    _smart_heading,
)
from validate_html import validate_html  # noqa: E402


SAMPLE_MD = """# 基准文章标题

这是第一段正文，用于确认中文字体、行高与段落节奏。原始标签 <img src="https://example.com/a.png"> 必须被转义。

## 第一部分

> 设计不是装饰，而是信息优先级的可见表达。

- 结构应当清晰
- 数据应当可核查
- 图片应当自包含

## 第二部分

| 指标 | 结果 | 说明 |
|---|---|---|
| 完整率 | 100% | 主题均通过验证 |
| 外链图片 | 0 | 发布版全部内嵌 |

### 实现细节

正文包含 **重点信息**、*术语* 和 `inline code`。

```python
print("quality")
```
"""


class ConverterTests(unittest.TestCase):
    def test_theme_set_and_legacy_aliases(self) -> None:
        self.assertEqual(len(THEMES), 10)
        self.assertEqual(canonical_theme("orange"), "vibrant")
        self.assertEqual(canonical_theme("nature"), "minimal")
        self.assertEqual(canonical_theme("blue"), "academic-blue")
        self.assertEqual(canonical_theme("morandi"), "fresh")
        self.assertEqual(set(ALIASES.values()), {"vibrant", "minimal", "academic-blue", "fresh"})

    def test_all_themes_render_and_validate(self) -> None:
        blocks = parse_markdown(SAMPLE_MD, ROOT)
        outputs = []
        for key in THEMES:
            document = render_html(blocks, "基准文章标题", key)
            self.assertEqual(validate_html(document, key), [], key)
            self.assertEqual(document.count("基准文章标题"), 2)  # title element and <title>
            self.assertIn("&lt;img src=&quot;https://example.com/a.png&quot;&gt;", document)
            outputs.append(document)
        self.assertEqual(len(set(outputs)), len(THEMES))

    def test_list_and_quote_close_before_following_blocks(self) -> None:
        blocks = parse_markdown("- item\n## heading\n> quote\nafter", ROOT)
        self.assertEqual([block.kind for block in blocks], ["list", "heading", "quote", "paragraph"])
        document = render_html(blocks, "T", "classic")
        self.assertNotIn("<ul", document[document.find("<h2"):])
        self.assertEqual(validate_html(document, "classic"), [])

    def test_external_markdown_image_is_rejected(self) -> None:
        with self.assertRaises(ValueError):
            parse_markdown("![bad](https://example.com/a.png)", ROOT)

    def test_unstyled_docx_editorial_structure_is_promoted(self) -> None:
        self.assertEqual(_smart_heading("第一道损耗：血清中的屏障"), 2)
        self.assertEqual(_smart_heading("三大工程策略：全链路耦合"), 2)
        blocks = _normalize_docx_blocks(
            [
                Block("paragraph", text="文章标题"),
                Block("paragraph", text="副标题：补充说明"),
                Block("quote", text="正文"),
                Block("image", data_uri="data:image/png;base64,AA=="),
                Block("paragraph", text="图 1：示意图"),
                Block("paragraph", text="来源：公开资料"),
                Block("paragraph", text="一句话：关键判断"),
            ]
        )
        self.assertEqual((blocks[0].kind, blocks[0].level), ("heading", 1))
        image = next(block for block in blocks if block.kind == "image")
        self.assertIn("来源：公开资料", image.caption)
        self.assertEqual(blocks[-1].kind, "quote")

    def test_preview_requires_explicit_validation_mode(self) -> None:
        blocks = parse_markdown("正文", ROOT)
        preview = render_html(blocks, "预览", "classic", preview_fonts=True, font_base="../assets/fonts")
        self.assertTrue(validate_html(preview, "classic"))
        self.assertEqual(validate_html(preview, "classic", allow_preview=True), [])

    def test_docx_parses_heading_table_and_image(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            image = root / "tiny.png"
            image.write_bytes(tiny_png)
            source = root / "sample.docx"
            doc = Document()
            doc.add_heading("Word 标题", level=1)
            doc.add_paragraph("正文段落")
            doc.add_picture(str(image))
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "指标"
            table.cell(0, 1).text = "结果"
            table.cell(1, 0).text = "图片"
            table.cell(1, 1).text = "已提取"
            doc.save(source)
            blocks, _ = parse_docx(source)
        kinds = [block.kind for block in blocks]
        self.assertIn("heading", kinds)
        self.assertIn("image", kinds)
        self.assertIn("table", kinds)
        image_block = next(block for block in blocks if block.kind == "image")
        self.assertTrue(image_block.data_uri.startswith("data:image/"))


if __name__ == "__main__":
    unittest.main()
